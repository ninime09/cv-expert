#!/usr/bin/env python3
"""
parse_resume.py — Parse a .docx resume into structured JSON.

Usage:
  python3 parse_resume.py \
    --input "/path/to/resume.docx" \
    --output "/tmp/resume_parsed.json"

Outputs:
  - Writes structured JSON to --output path
  - Prints a brief summary JSON to stdout for Claude to read

JSON structure:
  {
    "meta": { ... file info + ATS flags ... },
    "contact": { name, email, phone, linkedin, ... },
    "sections": [
      {
        "section_id": "sec_0",
        "heading": "Education",
        "heading_style": "Heading 1",
        "paragraphs": [
          {
            "para_id": "sec_0_p0",
            "raw_text": "...",      ← raw joined runs
            "full_text": "...",     ← normalized text (same normalization as write_resume.py)
            "style": "Normal",
            "is_bullet": false,
            "is_empty": false
          }
        ]
      }
    ],
    "flat_bullets": [
      {
        "para_id": "sec_1_p2",
        "section": "Work Experience",
        "raw_text": "...",
        "full_text": "..."
      }
    ],
    "ats_flags": {
      "has_tables": true,
      "textboxes_possible": false,    ← best-effort only
      "multi_column_possible": false, ← best-effort only
      "date_formats_found": ["June 2025", "Aug 2025 – Present"]
    }
  }
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

# Add scripts dir to path for text_utils import
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from text_utils import normalize_text, para_raw_text


# ── Section heading detection ─────────────────────────────────────────────────

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}

# Common resume section keywords used to detect sections even without heading styles
SECTION_KEYWORDS = [
    "education", "experience", "work experience", "internship", "projects",
    "skills", "technical skills", "awards", "honors", "publications",
    "certifications", "leadership", "activities", "volunteer", "summary",
    "objective", "profile", "contact", "languages", "coursework",
    "教育背景", "工作经历", "实习经历", "项目经历", "技能", "奖项",
    "荣誉", "证书", "领导力", "活动", "个人简介", "联系方式",
]

def is_section_heading(paragraph) -> bool:
    """Detect if a paragraph is a section heading."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name in HEADING_STYLES:
        return True
    # Heuristic: ALL CAPS or bold short line
    text = para_raw_text(paragraph).strip()
    if not text or len(text) > 80:
        return False
    if text.upper() == text and len(text) > 2:
        return True
    normalized_lower = text.lower()
    for kw in SECTION_KEYWORDS:
        if normalized_lower == kw or normalized_lower.startswith(kw):
            return True
    return False


# ── Bullet detection ──────────────────────────────────────────────────────────

BULLET_STYLE_PREFIXES = ("list", "bullet", "- ", "• ")

def is_bullet_paragraph(paragraph) -> bool:
    """Detect if a paragraph is a bullet point."""
    style_name = (paragraph.style.name or "").lower()
    if any(style_name.startswith(p) for p in BULLET_STYLE_PREFIXES):
        return True
    # Check for numPr (numbered/bulleted list) in XML
    try:
        pPr = paragraph._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        if pPr is not None:
            numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
            if numPr is not None:
                return True
    except Exception:
        pass
    # Text starts with common bullet characters
    text = para_raw_text(paragraph)
    if text and text[0] in ("•", "·", "‣", "▪", "▸", "○", "●", "-", "*", "✓", "✔", "→"):
        return True
    return False


# ── Contact info extraction ───────────────────────────────────────────────────

EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?[\d][\d\s\-().]{7,}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)
GITHUB_RE = re.compile(r"github\.com/[\w\-]+", re.IGNORECASE)

def extract_contact(paragraphs_text: list) -> dict:
    """Extract contact info from the first ~10 paragraphs."""
    contact = {
        "name": None,
        "email": None,
        "phone": None,
        "linkedin": None,
        "github": None,
        "other_links": [],
    }
    combined = "\n".join(paragraphs_text[:10])
    emails = EMAIL_RE.findall(combined)
    if emails:
        contact["email"] = emails[0]
    phones = PHONE_RE.findall(combined)
    if phones:
        contact["phone"] = phones[0].strip()
    linkedin = LINKEDIN_RE.search(combined)
    if linkedin:
        contact["linkedin"] = linkedin.group(0)
    github = GITHUB_RE.search(combined)
    if github:
        contact["github"] = github.group(0)
    # Heuristic: first non-empty short line is the name
    for line in paragraphs_text[:5]:
        line = line.strip()
        if line and len(line) < 50 and not EMAIL_RE.search(line) and not PHONE_RE.search(line):
            contact["name"] = line
            break
    return contact


# ── Date format detection ─────────────────────────────────────────────────────

DATE_PATTERNS = [
    re.compile(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}\b", re.IGNORECASE),
    re.compile(r"\b\d{4}\s*[.\-/]\s*\d{1,2}\b"),
    re.compile(r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\b", re.IGNORECASE),
    re.compile(r"\b\d{4}\b\s*[-–]\s*\b\d{4}\b"),
    re.compile(r"\b\d{1,2}/\d{4}\b"),
]

def find_date_formats(all_text: str) -> list:
    """Find all distinct date format patterns in the document."""
    found = set()
    for pat in DATE_PATTERNS:
        matches = pat.findall(all_text)
        for m in matches:
            if isinstance(m, str):
                found.add(m.strip())
    return sorted(list(found))


# ── ATS flags ─────────────────────────────────────────────────────────────────

def detect_ats_flags(doc) -> dict:
    """Detect ATS-problematic elements in the document."""
    has_tables = len(doc.tables) > 0

    # Textboxes — best-effort via XML namespace search
    textboxes_possible = False
    try:
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        txbx_elements = doc.element.findall(
            f".//{{{ns}}}txbxContent"
        )
        if txbx_elements:
            textboxes_possible = True
        # Also check drawing/inline shapes
        drawing_elements = doc.element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        )
        if drawing_elements:
            textboxes_possible = True
    except Exception:
        pass

    # Multi-column — best-effort via sectPr/cols
    multi_column_possible = False
    try:
        sectPr_elements = doc.element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
        )
        for sectPr in sectPr_elements:
            cols = sectPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cols"
            )
            if cols is not None:
                num = cols.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
                )
                if num and int(num) > 1:
                    multi_column_possible = True
                    break
    except Exception:
        pass

    return {
        "has_tables": has_tables,
        "textboxes_possible": textboxes_possible,
        "multi_column_possible": multi_column_possible,
        "date_formats_found": [],  # filled in later
    }


# ── Table content extraction ──────────────────────────────────────────────────

def extract_table_text(doc) -> list:
    """Extract text from tables as flat paragraphs (for sections inside tables)."""
    table_paras = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    raw = para_raw_text(para)
                    if raw.strip():
                        table_paras.append(raw)
    return table_paras


# ── Main parse logic ──────────────────────────────────────────────────────────

def parse_resume(input_path: str, output_path: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        print(json.dumps({
            "success": False,
            "error": "python-docx is not installed. Run: pip install python-docx",
        }), file=sys.stderr)
        sys.exit(2)

    try:
        doc = Document(input_path)
    except Exception as e:
        print(json.dumps({
            "success": False,
            "error": f"Failed to open document: {e}",
        }), file=sys.stderr)
        sys.exit(2)

    # Detect ATS flags
    ats_flags = detect_ats_flags(doc)

    # Extract table cell text for sections that live inside tables
    table_texts = extract_table_text(doc)
    has_table_content = len(table_texts) > 0

    # Main paragraph walk
    sections = []
    current_section = None
    current_section_idx = -1
    para_counter = {}  # section_id → count of paragraphs

    all_texts = []
    flat_bullets = []

    for para in doc.paragraphs:
        raw = para_raw_text(para)
        full = normalize_text(raw)
        all_texts.append(raw)

        # Skip completely empty paragraphs (but track them)
        if not full:
            continue

        if is_section_heading(para):
            current_section_idx += 1
            section_id = f"sec_{current_section_idx}"
            para_counter[section_id] = 0
            current_section = {
                "section_id": section_id,
                "heading": raw.strip(),
                "heading_style": para.style.name if para.style else "Unknown",
                "paragraphs": [],
            }
            sections.append(current_section)
        else:
            # If no section detected yet, create a default "Header/Contact" section
            if current_section is None:
                current_section_idx = 0
                section_id = "sec_0"
                para_counter[section_id] = 0
                current_section = {
                    "section_id": section_id,
                    "heading": "_header",
                    "heading_style": "auto-detected",
                    "paragraphs": [],
                }
                sections.append(current_section)

            section_id = current_section["section_id"]
            p_idx = para_counter[section_id]
            para_counter[section_id] += 1
            para_id = f"{section_id}_p{p_idx}"

            is_bullet = is_bullet_paragraph(para)
            para_obj = {
                "para_id": para_id,
                "raw_text": raw,
                "full_text": full,
                "style": para.style.name if para.style else "Unknown",
                "is_bullet": is_bullet,
                "is_empty": False,
            }
            current_section["paragraphs"].append(para_obj)

            if is_bullet:
                flat_bullets.append({
                    "para_id": para_id,
                    "section": current_section["heading"],
                    "raw_text": raw,
                    "full_text": full,
                })

    # Add table content as a special section if tables were found
    if table_texts:
        current_section_idx += 1
        table_section_id = f"sec_{current_section_idx}"
        table_section = {
            "section_id": table_section_id,
            "heading": "_table_content",
            "heading_style": "table",
            "paragraphs": [],
        }
        for i, text in enumerate(table_texts):
            full_t = normalize_text(text)
            para_obj = {
                "para_id": f"{table_section_id}_p{i}",
                "raw_text": text,
                "full_text": full_t,
                "style": "table-cell",
                "is_bullet": False,
                "is_empty": False,
            }
            table_section["paragraphs"].append(para_obj)
        sections.append(table_section)

    # Date format detection
    all_text_combined = "\n".join(all_texts) + "\n" + "\n".join(table_texts)
    ats_flags["date_formats_found"] = find_date_formats(all_text_combined)

    # Contact extraction (from first section paragraphs)
    early_texts = [p["raw_text"] for s in sections[:2] for p in s["paragraphs"]]
    contact = extract_contact(early_texts)

    # Warnings
    warnings = []
    if ats_flags["has_tables"]:
        warnings.append(
            "Tables detected. Some ATS systems cannot parse table content. "
            "Consider converting table sections to plain text."
        )
    if ats_flags["textboxes_possible"]:
        warnings.append(
            "Possible text boxes or drawing objects detected (best-effort). "
            "ATS systems often skip text box content."
        )
    if ats_flags["multi_column_possible"]:
        warnings.append(
            "Multi-column layout possibly detected (best-effort). "
            "ATS systems may misread column order."
        )
    if len(sections) <= 1:
        warnings.append(
            "Very few sections detected. The resume structure may not follow "
            "standard heading conventions — review the parsed output."
        )

    result = {
        "meta": {
            "source_file": str(input_path),
            "output_file": str(output_path),
            "parse_timestamp": datetime.now(timezone.utc).isoformat(),
            "parser_version": "1.0.0",
            "total_paragraphs": sum(len(s["paragraphs"]) for s in sections),
            "total_sections": len(sections),
            "total_bullets": len(flat_bullets),
            "warnings": warnings,
        },
        "contact": contact,
        "sections": sections,
        "flat_bullets": flat_bullets,
        "ats_flags": ats_flags,
    }

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    return result


def main():
    parser = argparse.ArgumentParser(description="Parse a .docx resume to structured JSON.")
    parser.add_argument("--input", required=True, help="Path to input .docx file")
    parser.add_argument("--output", required=True, help="Path to write output JSON")
    args = parser.parse_args()

    result = parse_resume(args.input, args.output)

    # Print concise summary to stdout for Claude to read
    summary = {
        "success": True,
        "output_file": args.output,
        "meta": result["meta"],
        "contact_name": result["contact"].get("name"),
        "sections_found": [s["heading"] for s in result["sections"]],
        "total_bullets": result["meta"]["total_bullets"],
        "ats_flags": result["ats_flags"],
        "warnings": result["meta"]["warnings"],
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
