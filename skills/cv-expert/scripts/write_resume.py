#!/usr/bin/env python3
"""
write_resume.py — Apply patches to a .docx resume and add Word comments.

Usage:
  python3 write_resume.py \
    --input  "/path/to/resume.docx" \
    --patches "/path/to/patches.json" \
    --output  "/path/to/resume_optimized.docx" \
    --backup  "/path/to/resume_backup.docx" \
    [--author "Resume Optimizer AI"] \
    [--initials "RO"]

V1 supported patch types:
  - replace_text: Replace entire paragraph text + add comment
  - comment_only: Add comment to paragraph without text change

Output JSON to stdout:
  {
    "success": true,
    "output_file": "...",
    "backup_file": "...",
    "output_dir_used": "original|downloads|tmp",
    "patches_applied": N,
    "patches_failed": N,
    "failed_patch_ids": [...],
    "comments_added": N
  }
"""

import argparse
import json
import os
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path

# Add scripts dir to path for text_utils import
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from text_utils import normalize_text, para_raw_text, fuzzy_match, contains_match, exact_match


# ── Output path resolution with fallback ─────────────────────────────────────

def resolve_output_path(output_path: str, backup_path: str) -> tuple:
    """
    Resolve output and backup paths with fallback.
    Priority: requested path → ~/Downloads → /tmp

    Returns (output_path, backup_path, dir_used: str)
    """
    def try_dir(dirpath: str, filename: str) -> str | None:
        full = os.path.join(dirpath, filename)
        try:
            # Test writability with a dummy write
            test = full + ".writetest"
            with open(test, "w") as f:
                f.write("")
            os.remove(test)
            return full
        except Exception:
            return None

    out_dir = os.path.dirname(os.path.abspath(output_path))
    out_file = os.path.basename(output_path)
    bak_file = os.path.basename(backup_path)

    for directory, label in [
        (out_dir, "original"),
        (str(Path.home() / "Downloads"), "downloads"),
        ("/tmp", "tmp"),
    ]:
        resolved_out = try_dir(directory, out_file)
        resolved_bak = try_dir(directory, bak_file)
        if resolved_out and resolved_bak:
            return resolved_out, resolved_bak, label

    # Last resort: use /tmp with flat filenames
    return (
        f"/tmp/{out_file}",
        f"/tmp/{bak_file}",
        "tmp",
    )


# ── Paragraph lookup ──────────────────────────────────────────────────────────

HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
SECTION_KEYWORDS = [
    "education", "experience", "work experience", "internship", "projects",
    "skills", "technical skills", "awards", "honors", "publications",
    "certifications", "leadership", "activities", "volunteer", "summary",
    "objective", "profile", "contact", "languages", "coursework",
    "教育背景", "工作经历", "实习经历", "项目经历", "技能", "奖项",
    "荣誉", "证书", "领导力", "活动", "个人简介", "联系方式",
]

def _is_section_heading(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name in HEADING_STYLES:
        return True
    text = para_raw_text(paragraph).strip()
    if not text or len(text) > 80:
        return False
    if text.upper() == text and len(text) > 2:
        return True
    normalized_lower = text.lower()
    for kw in SECTION_KEYWORDS:
        if normalized_lower == kw or normalized_lower.startswith(kw):
            return True
    return False


def build_para_index(doc) -> dict:
    """
    Build a deterministic {para_id: paragraph_object} index.
    Must mirror the same logic as parse_resume.py.
    """
    index = {}
    current_section_idx = -1
    current_section_id = None
    para_counter = {}

    for para in doc.paragraphs:
        raw = para_raw_text(para)
        full = normalize_text(raw)

        if not full:
            continue

        if _is_section_heading(para):
            current_section_idx += 1
            current_section_id = f"sec_{current_section_idx}"
            para_counter[current_section_id] = 0
        else:
            if current_section_id is None:
                current_section_idx = 0
                current_section_id = "sec_0"
                para_counter[current_section_id] = 0

            p_idx = para_counter[current_section_id]
            para_counter[current_section_id] += 1
            para_id = f"{current_section_id}_p{p_idx}"
            index[para_id] = para

    return index


# ── Word comment addition ─────────────────────────────────────────────────────

def add_comment(doc, paragraph, comment_text: str, author: str, initials: str):
    """
    Add a Word comment anchored to a paragraph.

    python-docx 1.2.0+ supports doc.add_comment(runs, text, author, initials).
    For older versions, we fall back to a manual XML approach.
    """
    runs = paragraph.runs
    if not runs:
        # Paragraph has no runs (e.g., empty), create a dummy run to anchor
        run = paragraph.add_run("")
        runs = [run]

    try:
        # python-docx >= 1.2.0 native API
        doc.add_comment(
            runs=runs,
            text=comment_text,
            author=author,
            initials=initials,
        )
    except AttributeError:
        # Fallback: inject comment XML manually
        _add_comment_xml(doc, paragraph, comment_text, author, initials)


def _add_comment_xml(doc, paragraph, comment_text: str, author: str, initials: str):
    """
    Manual XML-based comment insertion for python-docx < 1.2.0.
    This is a best-effort fallback and may not render in all Word viewers.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    # Get or create comments part
    comments_part = None
    try:
        comments_part = doc.part.comments_part
    except Exception:
        pass

    if comments_part is None:
        return  # Cannot add comments without comments_part

    comments_el = comments_part._element
    # Assign a unique comment ID
    existing = comments_el.findall(qn("w:comment"))
    comment_id = str(len(existing))

    # Build the comment element
    comment_el = OxmlElement("w:comment")
    comment_el.set(qn("w:id"), comment_id)
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    comment_el.set(qn("w:initials"), initials)

    comment_para = OxmlElement("w:p")
    comment_run = OxmlElement("w:r")
    comment_text_el = OxmlElement("w:t")
    comment_text_el.text = comment_text
    comment_text_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    comment_run.append(comment_text_el)
    comment_para.append(comment_run)
    comment_el.append(comment_para)
    comments_el.append(comment_el)

    # Add comment reference marks around the paragraph runs
    para_el = paragraph._p
    first_run = para_el.find(qn("w:r"))
    if first_run is None:
        return

    rangeStart = OxmlElement("w:commentRangeStart")
    rangeStart.set(qn("w:id"), comment_id)
    rangeEnd = OxmlElement("w:commentRangeEnd")
    rangeEnd.set(qn("w:id"), comment_id)
    ref = OxmlElement("w:r")
    ref_el = OxmlElement("w:commentReference")
    ref_el.set(qn("w:id"), comment_id)
    ref.append(ref_el)

    para_el.insert(list(para_el).index(first_run), rangeStart)
    para_el.append(rangeEnd)
    para_el.append(ref)


# ── Patch application ─────────────────────────────────────────────────────────

def match_paragraph(para, original_text: str, match_mode: str) -> bool:
    """Check if a paragraph matches original_text using the specified match_mode."""
    full = normalize_text(para_raw_text(para))
    if match_mode == "exact":
        return exact_match(normalize_text(original_text), full)
    elif match_mode == "fuzzy":
        return fuzzy_match(original_text, full)
    else:  # contains (default)
        return normalize_text(original_text) in full


def apply_replace_text(doc, para, patch: dict, author: str, initials: str):
    """
    V1 replace_text: Replace the ENTIRE paragraph text with replacement_text.
    Clears all existing runs and writes a single new run, preserving paragraph style.
    """
    # Remove all existing runs
    p_el = para._p
    from docx.oxml.ns import qn
    for r in p_el.findall(qn("w:r")):
        p_el.remove(r)

    # Add new run with replacement text
    new_run = para.add_run(patch["replacement_text"])

    # Add comment
    add_comment(doc, para, patch["comment"], author, initials)
    return True


def apply_comment_only(doc, para, patch: dict, author: str, initials: str):
    """Add a Word comment to a paragraph without changing any text."""
    add_comment(doc, para, patch["comment"], author, initials)
    return True


def apply_patches(doc, patches: list, para_index: dict, author: str, initials: str) -> dict:
    """Apply all patches to the document. Returns a result summary."""
    applied = 0
    failed = 0
    failed_ids = []
    comments_added = 0

    for patch in patches:
        patch_id = patch.get("patch_id", "unknown")
        patch_type = patch.get("type", "")
        para_id = patch.get("target_para_id", "")
        original_text = patch.get("original_text", "")
        match_mode = patch.get("match_mode", "contains")
        occurrence_index = patch.get("occurrence_index", 0)

        # Find target paragraph by para_id
        para = para_index.get(para_id)
        if para is None:
            print(
                f"[write_resume] WARNING: para_id '{para_id}' not found in document (patch {patch_id})",
                file=sys.stderr,
            )
            failed += 1
            failed_ids.append(patch_id)
            continue

        # Verify original_text match (for replace_text)
        if patch_type == "replace_text" and original_text:
            if not match_paragraph(para, original_text, match_mode):
                print(
                    f"[write_resume] WARNING: original_text not matched in para '{para_id}' "
                    f"(match_mode={match_mode}, patch {patch_id}). Applying comment-only fallback.",
                    file=sys.stderr,
                )
                # Fallback: add comment explaining the mismatch
                fallback_comment = (
                    f"[AUTO-FALLBACK] Could not match original text for replacement. "
                    f"Intended change: {patch.get('replacement_text', '')} | "
                    f"Original comment: {patch.get('comment', '')}"
                )
                try:
                    apply_comment_only(doc, para, {"comment": fallback_comment}, author, initials)
                    comments_added += 1
                    applied += 1
                except Exception as e:
                    failed += 1
                    failed_ids.append(patch_id)
                continue

        try:
            if patch_type == "replace_text":
                success = apply_replace_text(doc, para, patch, author, initials)
            elif patch_type == "comment_only":
                success = apply_comment_only(doc, para, patch, author, initials)
            else:
                print(
                    f"[write_resume] WARNING: Unknown patch type '{patch_type}' (patch {patch_id})",
                    file=sys.stderr,
                )
                failed += 1
                failed_ids.append(patch_id)
                continue

            if success:
                applied += 1
                comments_added += 1
            else:
                failed += 1
                failed_ids.append(patch_id)

        except Exception as e:
            print(
                f"[write_resume] ERROR applying patch {patch_id}: {e}",
                file=sys.stderr,
            )
            failed += 1
            failed_ids.append(patch_id)

    return {
        "patches_applied": applied,
        "patches_failed": failed,
        "failed_patch_ids": failed_ids,
        "comments_added": comments_added,
    }


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Apply patches to a .docx resume.")
    parser.add_argument("--input", required=True, help="Path to original .docx")
    parser.add_argument("--patches", required=True, help="Path to patches.json")
    parser.add_argument("--output", required=True, help="Path for optimized .docx output")
    parser.add_argument("--backup", required=True, help="Path for backup .docx")
    parser.add_argument("--author", default="Resume Optimizer AI", help="Comment author name")
    parser.add_argument("--initials", default="RO", help="Comment author initials")
    args = parser.parse_args()

    try:
        from docx import Document
    except ImportError:
        print(json.dumps({
            "success": False,
            "error": "python-docx is not installed. Run: pip install python-docx",
        }))
        sys.exit(2)

    # Load patches
    try:
        with open(args.patches, "r", encoding="utf-8") as f:
            patches_data = json.load(f)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"Failed to load patches: {e}"}))
        sys.exit(1)

    patches = patches_data.get("patches", [])
    if not patches:
        print(json.dumps({"success": False, "error": "patches.json contains no patches."}))
        sys.exit(1)

    # Resolve output paths with fallback
    output_path, backup_path, dir_used = resolve_output_path(args.output, args.backup)

    # Create backup FIRST before any modification
    try:
        shutil.copy2(args.input, backup_path)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"Failed to create backup: {e}"}))
        sys.exit(1)

    # Open document for modification
    try:
        doc = Document(args.input)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"Failed to open document: {e}"}))
        sys.exit(1)

    # Build paragraph index
    para_index = build_para_index(doc)

    # Apply patches
    result = apply_patches(doc, patches, para_index, args.author, args.initials)

    # Save optimized document
    try:
        doc.save(output_path)
    except Exception as e:
        print(json.dumps({"success": False, "error": f"Failed to save output: {e}"}))
        sys.exit(1)

    output = {
        "success": True,
        "output_file": output_path,
        "backup_file": backup_path,
        "output_dir_used": dir_used,
        **result,
    }
    print(json.dumps(output, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
